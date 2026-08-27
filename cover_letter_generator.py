"""
cover_letter_generator.py — Gemini-powered Job-Specific Cover Letter Generator.

Creates a compelling, recruiter-targeting 3-paragraph cover letter:
1. High-Impact Hook: Tailored to the company and exact role title.
2. Core Accomplishments: Connects candidate's top projects/experience to the JD requirements & missing keywords.
3. Call to Action: Professional closing requesting an interview.
"""

import json
import os
import re
from pathlib import Path
from google import genai
from google.genai import types
from gemini_client import get_gemini_client, execute_with_failover, get_all_gemini_keys


def _get_gemini_client():
    return get_gemini_client()


def generate_cover_letter(
    base_resume: dict,
    jd_text: str,
    company: str,
    role: str,
    missing_keywords: list = None,
    output_dir: str = None,
) -> dict:
    """
    Generate a high-impact, job-specific cover letter using Gemini.

    Returns:
        dict with keys: {
            "cover_letter_text": str,
            "file_path_docx": str,
            "file_path_txt": str,
            "company": str,
            "role": str
        }
    """
    missing_keywords = missing_keywords or []
    client = _get_gemini_client()

    candidate_name = base_resume.get("name", "Applicant")
    candidate_email = base_resume.get("contact", {}).get("email", "")
    candidate_phone = base_resume.get("contact", {}).get("phone", "")
    candidate_location = base_resume.get("contact", {}).get("location", "")

    prompt = f"""You are an elite Executive Career Strategist and Resume Architect.
Write a compelling, recruiter-targeting 3-paragraph Cover Letter for {candidate_name} applying for the {role} position at {company}.

CANDIDATE PROFILE:
Name: {candidate_name}
Email: {candidate_email} | Phone: {candidate_phone} | Location: {candidate_location}
Summary: {base_resume.get('summary', '')}
Top Skills: {", ".join(base_resume.get('skills', [])[:15])}
Recent Experience: {json.dumps(base_resume.get('experience', [])[:2], indent=2)}

TARGET JOB DETAILS:
Company: {company}
Role: {role}
Key Terms to Weave In: {", ".join(missing_keywords[:10]) if missing_keywords else "N/A"}

JOB DESCRIPTION:
{jd_text[:3500]}

COVER LETTER STRUCTURAL REQUIREMENTS:
- Length: Exactly 3 structured, powerful paragraphs (250 - 350 words total).
- Paragraph 1 (The Hook): Express genuine enthusiasm for {role} at {company}. State why candidate's experience in full-stack engineering and cloud architecture aligns perfectly with their mission.
- Paragraph 2 (The Proof & Impact): Highlight 2-3 specific technical achievements from experience that match the key requirements of the job description. Weave in key technical terms smoothly.
- Paragraph 3 (The Closing): Express confidence in driving results for {company}. Request an interview.

TONE & STYLE RULES:
- Professional, confident, active human voice.
- NO cliché AI buzzwords ("spearheaded", "leveraged", "pivotal", "testament", "transformative", "groundbreaking", "fostered").
- Return ONLY the clean cover letter text including date and recipient header. No markdown code blocks around the text.
"""

    def _call_gemini(client):
        for model in ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-flash"]:
            try:
                response = client.models.generate_content(
                    model=model,
                    contents=[
                        types.Content(
                            role="user",
                            parts=[types.Part(text=prompt)],
                        )
                    ],
                    config=types.GenerateContentConfig(
                        temperature=0.4,
                        max_output_tokens=1500,
                    ),
                )
                if response.text and len(response.text.strip()) > 100:
                    print(f"[Cover Letter] Generated successfully with {model} ({len(response.text.strip())} chars)")
                    return response.text.strip()
            except Exception as e:
                print(f"[Cover Letter] Model {model} note: {e}")
                if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                    raise e
        return ""

    try:
        cover_letter_text = execute_with_failover(_call_gemini)
    except Exception as e:
        print(f"[Cover Letter] Generation failed across all keys: {e}")
        cover_letter_text = ""

    if not cover_letter_text:
        # Fallback template if all Gemini calls fail
        cover_letter_text = f"""Dear Hiring Team at {company},

I am writing to express my strong enthusiasm for the {role} position. With extensive experience architecting high-scalability web applications and SaaS platforms, I am confident in my ability to deliver immediate value to {company}.

Throughout my career, I have engineered robust frontend and backend microservices, optimized database performance, and led cross-functional teams to deliver critical product features on time. My experience aligns closely with your core requirements, particularly in scalable software architecture and full-stack development.

I would welcome the opportunity to discuss how my background and technical capabilities can support {company}'s goals. Thank you for your time and consideration.

Sincerely,
{candidate_name}
{candidate_email} | {candidate_phone}"""

    # ── Save .docx and .txt files in output folder ───────────────────────────
    file_path_docx = ""
    file_path_txt = ""

    if output_dir:
        from resume_builder import slugify
        folder_path = Path(output_dir)
        folder_path.mkdir(parents=True, exist_ok=True)

        name_slug = slugify(candidate_name)
        out_name = f"{name_slug}_Cover_Letter"

        file_path_txt = str(folder_path / f"{out_name}.txt")
        file_path_docx = str(folder_path / f"{out_name}.docx")

        # Write .txt file
        with open(file_path_txt, "w", encoding="utf-8") as f:
            f.write(cover_letter_text)

        # Write .docx file
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            for s in doc.sections:
                s.top_margin = Inches(0.8)
                s.bottom_margin = Inches(0.8)
                s.left_margin = Inches(0.8)
                s.right_margin = Inches(0.8)

            paragraphs = cover_letter_text.split("\n\n")
            for i, p_text in enumerate(paragraphs):
                p_text = p_text.strip()
                if not p_text:
                    continue
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(8)
                para.paragraph_format.line_spacing = 1.15
                run = para.add_run(p_text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                if i == 0 and ("Dear" in p_text or candidate_name in p_text):
                    run.font.bold = True

            doc.save(file_path_docx)
            print(f"[Cover Letter] Saved DOCX: {file_path_docx}")
        except Exception as docx_err:
            print(f"[Cover Letter] DOCX save note: {docx_err}")

    return {
        "cover_letter_text": cover_letter_text,
        "file_path_docx": file_path_docx,
        "file_path_txt": file_path_txt,
        "company": company,
        "role": role,
    }
