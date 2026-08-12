"""
resume_builder.py — Generates a clean, ATS-optimized Word document from the final resume JSON.
Single font, clean bullets, no special characters that break ATS parsing.
"""
import json
import os
import re
from pathlib import Path
from datetime import datetime


def slugify(text: str) -> str:
    """Convert text to a safe folder/file name."""
    text = re.sub(r'[^\w\s-]', '', str(text))
    text = re.sub(r'[\s_-]+', '_', text)
    return text.strip('_')


def sanitize_text(text: str) -> str:
    """Remove/replace characters that break ATS parsing."""
    if not text:
        return ""
    # Replace em dashes and en dashes with hyphens
    text = text.replace('\u2014', '-').replace('\u2013', '-')
    # Replace fancy quotes
    text = text.replace('\u2019', "'").replace('\u2018', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    # Remove bullet characters (we'll use docx bullets instead)
    text = re.sub(r'^[•·▪▸►\*]\s*', '', text)
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def build_resume_docx(
    resume: dict,
    company: str,
    role: str,
    output_dir: str = None,
) -> str:
    """
    Build a clean Word document from the resume dict.
    
    Args:
        resume: Final cleaned resume dict
        company: Company name (for folder/file naming)
        role: Job role (for folder/file naming)
        output_dir: Base output directory (default: job-agent/output/)
    
    Returns:
        Absolute path to the generated .docx file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    
    # ─── Setup output path ───────────────────────────────────────────────────
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(__file__), "output")
    
    folder_name = f"{slugify(company)}_{slugify(role)}"[:80]
    folder_path = Path(output_dir) / folder_name
    folder_path.mkdir(parents=True, exist_ok=True)
    
    file_path = folder_path / "Jalal_Khan_Resume.docx"
    
    # ─── Create Document ─────────────────────────────────────────────────────
    doc = Document()
    
    # Set page margins (narrow for more content space)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ─── Helper functions ─────────────────────────────────────────────────────
    
    def set_font(run, size=10.5, bold=False, italic=False, color=None):
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    
    def add_paragraph_spacing(para, space_before=0, space_after=0):
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)
    
    def add_section_header(title: str):
        """Add a formatted section header with a bottom border line."""
        para = doc.add_paragraph()
        add_paragraph_spacing(para, space_before=8, space_after=2)
        run = para.add_run(title.upper())
        set_font(run, size=10.5, bold=True, color=(31, 73, 125))  # Dark blue
        
        # Add bottom border to paragraph
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        return para
    
    def add_bullet(text: str):
        """Add a clean bullet point."""
        text = sanitize_text(text)
        if not text:
            return
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(text)
        set_font(run, size=10)
        add_paragraph_spacing(para, space_before=1, space_after=1)
        para.paragraph_format.left_indent = Inches(0.25)
    
    def add_normal_text(text: str, size=10, bold=False, italic=False):
        """Add a normal paragraph."""
        text = sanitize_text(text)
        para = doc.add_paragraph()
        run = para.add_run(text)
        set_font(run, size=size, bold=bold)
        run.italic = italic
        add_paragraph_spacing(para, space_before=1, space_after=1)
        return para
    
    def add_two_column_line(left: str, right: str, left_bold=False, right_italic=False):
        """Add a line with text on left and right (e.g., role + dates)."""
        para = doc.add_paragraph()
        add_paragraph_spacing(para, space_before=3, space_after=1)
        
        run_left = para.add_run(sanitize_text(left))
        set_font(run_left, size=10.5, bold=left_bold)
        
        if right:
            # Tab to right-align the date
            tab_stop = OxmlElement('w:tab')
            run_left._r.append(tab_stop)
            
            # Add right-aligned tab stop at page right margin
            pPr = para._p.get_or_add_pPr()
            tabs = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '9360')  # ~6.5 inches in twips
            tabs.append(tab)
            pPr.append(tabs)
            
            run_right = para.add_run(sanitize_text(right))
            set_font(run_right, size=10, italic=right_italic)
        
        return para
    
    # ─── HEADER: Name ────────────────────────────────────────────────────────
    contact = resume.get("contact", {})
    name = sanitize_text(resume.get("name", "Jalal Khan"))
    
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_spacing(name_para, space_before=0, space_after=3)
    run = name_para.add_run(name)
    set_font(run, size=18, bold=True, color=(31, 73, 125))
    
    # ─── HEADER: Contact line ─────────────────────────────────────────────────
    contact_parts = []
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("linkedin"):
        contact_parts.append(contact["linkedin"])
    if contact.get("github"):
        contact_parts.append(contact["github"])
    if contact.get("portfolio"):
        contact_parts.append(contact["portfolio"])
    if contact.get("location"):
        contact_parts.append(contact["location"])
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph_spacing(contact_para, space_before=0, space_after=6)
        run = contact_para.add_run("  |  ".join(contact_parts))
        set_font(run, size=9, color=(89, 89, 89))
    
    # ─── SUMMARY ─────────────────────────────────────────────────────────────
    summary = resume.get("summary", "")
    if summary:
        add_section_header("Professional Summary")
        add_normal_text(sanitize_text(summary), size=10)
    
    # ─── SKILLS ──────────────────────────────────────────────────────────────
    skills = resume.get("skills", [])
    if skills:
        add_section_header("Technical Skills")
        
        # Group skills into rows of ~6 per line for clean layout
        clean_skills = [sanitize_text(s) for s in skills if s]
        
        # Try to categorize if possible
        categories = _categorize_skills(clean_skills)
        
        for category, skill_list in categories.items():
            if not skill_list:
                continue
            para = doc.add_paragraph()
            add_paragraph_spacing(para, space_before=1, space_after=1)
            
            label_run = para.add_run(f"{category}: ")
            set_font(label_run, size=10, bold=True)
            
            skills_run = para.add_run(", ".join(skill_list))
            set_font(skills_run, size=10)
    
    # ─── EXPERIENCE ──────────────────────────────────────────────────────────
    experience = resume.get("experience", [])
    if experience:
        add_section_header("Professional Experience")
        
        for role_entry in experience:
            title = sanitize_text(role_entry.get("title", ""))
            company_name = sanitize_text(role_entry.get("company", ""))
            dates = sanitize_text(role_entry.get("dates", ""))
            location = sanitize_text(role_entry.get("location", ""))
            bullets = role_entry.get("bullets", [])
            
            if not title:
                continue
            
            # Role title + dates on same line
            title_company = f"{title} - {company_name}" if company_name else title
            add_two_column_line(title_company, dates, left_bold=True, right_italic=True)
            
            # Location (if available)
            if location:
                add_normal_text(location, size=9.5, italic=True)
            
            # Bullet points
            for bullet in bullets:
                if bullet and len(sanitize_text(bullet)) > 5:
                    add_bullet(bullet)
            
            # Small spacing between roles
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
    
    # ─── PROJECTS ────────────────────────────────────────────────────────────
    projects = resume.get("projects", [])
    if projects:
        add_section_header("Projects")
        
        for project in projects:
            proj_name = sanitize_text(project.get("name", ""))
            tech_stack = project.get("tech_stack", [])
            description = sanitize_text(project.get("description", ""))
            url = sanitize_text(project.get("url", ""))
            
            if not proj_name:
                continue
            
            # Project name + tech stack
            tech_str = f" | {', '.join(tech_stack)}" if tech_stack else ""
            url_str = f" | {url}" if url else ""
            
            para = doc.add_paragraph()
            add_paragraph_spacing(para, space_before=3, space_after=1)
            run = para.add_run(f"{proj_name}{tech_str}{url_str}")
            set_font(run, size=10.5, bold=True)
            
            if description:
                add_bullet(description)
    
    # ─── EDUCATION ───────────────────────────────────────────────────────────
    education = resume.get("education", [])
    if education:
        add_section_header("Education")
        
        for edu in education:
            institution = sanitize_text(edu.get("institution", ""))
            degree = sanitize_text(edu.get("degree", ""))
            field = sanitize_text(edu.get("field", ""))
            grad_date = sanitize_text(edu.get("graduation_date", ""))
            gpa = sanitize_text(edu.get("gpa", ""))
            
            if not institution and not degree:
                continue
            
            degree_field = f"{degree}" + (f" in {field}" if field else "")
            add_two_column_line(institution, grad_date, left_bold=True, right_italic=True)
            
            if degree_field:
                add_normal_text(degree_field, size=10, italic=True)
            if gpa:
                add_normal_text(f"GPA: {gpa}", size=9.5)
    
    # ─── CERTIFICATIONS ──────────────────────────────────────────────────────
    certifications = resume.get("certifications", [])
    if certifications:
        add_section_header("Certifications")
        for cert in certifications:
            if cert:
                add_bullet(sanitize_text(cert))
    
    # ─── Save ─────────────────────────────────────────────────────────────────
    doc.save(str(file_path))
    
    print(f"[Builder] [OK] Resume saved: {file_path}")
    print(f"[Builder]    Folder: {folder_path}")
    
    return str(file_path)


def _categorize_skills(skills: list) -> dict:
    """Auto-categorize skills into groups for clean display."""
    categories = {
        "Languages": [],
        "Frameworks & Libraries": [],
        "Databases": [],
        "Cloud & DevOps": [],
        "Tools & Platforms": [],
        "Other": [],
    }
    
    lang_keywords = {
        'python', 'javascript', 'typescript', 'java', 'kotlin', 'swift', 'go', 'golang',
        'rust', 'c++', 'c#', 'ruby', 'php', 'scala', 'r', 'sql', 'html', 'css', 'bash', 'shell'
    }
    framework_keywords = {
        'react', 'vue', 'angular', 'next', 'nuxt', 'svelte', 'node', 'express', 'fastapi',
        'django', 'flask', 'spring', 'rails', 'laravel', 'graphql', 'rest', 'grpc',
        'tailwind', 'bootstrap', 'redux', 'react native', 'expo', 'pytorch', 'tensorflow'
    }
    db_keywords = {
        'postgresql', 'postgres', 'mysql', 'mongodb', 'sqlite', 'redis', 'elasticsearch',
        'dynamodb', 'cassandra', 'neo4j', 'supabase', 'firebase', 'pinecone', 'snowflake'
    }
    cloud_keywords = {
        'aws', 'gcp', 'azure', 'docker', 'kubernetes', 'terraform', 'ansible', 'ci/cd',
        'github actions', 'jenkins', 'vercel', 'netlify', 'heroku', 'cloudflare',
        'lambda', 'ec2', 's3', 'cloud run', 'gke', 'ecs'
    }
    
    for skill in skills:
        skill_lower = skill.lower()
        
        if any(kw in skill_lower for kw in lang_keywords):
            categories["Languages"].append(skill)
        elif any(kw in skill_lower for kw in framework_keywords):
            categories["Frameworks & Libraries"].append(skill)
        elif any(kw in skill_lower for kw in db_keywords):
            categories["Databases"].append(skill)
        elif any(kw in skill_lower for kw in cloud_keywords):
            categories["Cloud & DevOps"].append(skill)
        elif len(skill) > 2:
            categories["Tools & Platforms"].append(skill)
    
    # Remove empty categories
    return {k: v for k, v in categories.items() if v}


def convert_to_pdf(doc_path: str) -> str:
    """Convert .docx file to .pdf using Word COM on Windows or docx2pdf."""
    if not doc_path or not os.path.exists(doc_path):
        return ""

    pdf_path = str(Path(doc_path).with_suffix(".pdf"))
    if os.path.exists(pdf_path):
        return pdf_path

    # Attempt 1: Native Windows Word COM Automation (Fastest, Pixel-Perfect)
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 = wdFormatPDF
        doc.Close()
        word.Quit()
        if os.path.exists(pdf_path):
            print(f"[PDF] [OK] Converted via Word COM: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"[PDF] Word COM conversion note: {e}")

    # Attempt 2: docx2pdf fallback
    try:
        from docx2pdf import convert
        convert(doc_path, pdf_path)
        if os.path.exists(pdf_path):
            print(f"[PDF] [OK] Converted via docx2pdf: {pdf_path}")
            return pdf_path
    except Exception as e:
        print(f"[PDF] docx2pdf note: {e}")

    return pdf_path


if __name__ == "__main__":
    """Quick test of the Word builder."""
    import sys
    
    resume_path = os.path.join(os.path.dirname(__file__), "base_resume.json")
    if not os.path.exists(resume_path):
        print("ERROR: base_resume.json not found.")
        sys.exit(1)
    
    with open(resume_path, 'r', encoding='utf-8') as f:
        resume = json.load(f)
    
    output = build_resume_docx(resume, "TestCompany", "Software Engineer")
    print(f"Generated: {output}")
