"""
docx_patcher.py - In-place DOCX editor that preserves original styling.

Instead of building a new DOCX from scratch, this module:
1. Loads the user's original formatted DOCX (Canva / Word template)
2. Finds the old experience bullet text using fuzzy matching
3. Replaces ONLY the text content while keeping all Run formatting intact
   (fonts, colors, sizes, bold/italic, spacing, column layouts, etc.)
4. Saves as a new file in the output folder

This way the final resume keeps 100% of the original Canva design.
"""
import json
import os
import re
import shutil
from pathlib import Path
from datetime import datetime


def slugify(text: str) -> str:
    text = re.sub(r'[^\w\s-]', '', str(text))
    text = re.sub(r'[\s_-]+', '_', text)
    return text.strip('_')


def _get_para_full_text(para) -> str:
    """Get full text of a paragraph across all runs."""
    return "".join(run.text for run in para.runs)


def _set_para_text_preserve_format(para, new_text: str):
    """
    Replace paragraph text while preserving the formatting of the first run.
    Clears all extra runs, keeps formatting of run[0].
    """
    if not para.runs:
        para.add_run(new_text)
        return

    # Keep the first run's formatting, put all new text in it
    first_run = para.runs[0]
    first_run.text = new_text

    # Clear all other runs
    for run in para.runs[1:]:
        run.text = ""


def _normalize(text: str) -> str:
    """Normalize text for fuzzy comparison."""
    return re.sub(r'\s+', ' ', text.lower().strip())


def _fuzzy_match_score(a: str, b: str) -> float:
    """
    Simple word-overlap ratio between two strings.
    Returns 0.0 - 1.0 where 1.0 is identical.
    """
    a_words = set(_normalize(a).split())
    b_words = set(_normalize(b).split())
    if not a_words or not b_words:
        return 0.0
    intersection = a_words & b_words
    return len(intersection) / max(len(a_words), len(b_words))


def _get_all_paragraphs_from_doc(doc):
    """
    Yield all paragraphs from both the body AND table cells, in document order.
    Critical for DOCX files with table-based layouts (common in Canva exports).
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def iter_blocks(parent):
        for child in parent.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                yield Paragraph(child, parent)
            elif tag == 'tbl':
                yield from _iter_table_paragraphs(Table(child, parent))

    yield from iter_blocks(doc)


def _iter_table_paragraphs(table):
    """Recursively yield all paragraphs in all table cells."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para
            # Nested tables
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def patch_docx_with_rewritten_resume(
    original_docx_path: str,
    rewritten_resume: dict,
    company: str,
    role: str,
    output_dir: str = None,
) -> str:
    """
    Patch the original DOCX with rewritten content, preserving all formatting.

    Args:
        original_docx_path: Path to the user's original uploaded DOCX
        rewritten_resume: The Gemini-rewritten resume dict (from rewriter.py)
        company: For output folder naming
        role: For output folder naming
        output_dir: Base output directory

    Returns:
        Path to the patched DOCX file
    """
    from docx import Document

    if not os.path.exists(original_docx_path):
        raise FileNotFoundError(f"Original DOCX not found: {original_docx_path}")

    # Set up output path
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(__file__), "output")

    folder_name = f"{slugify(company)}_{slugify(role)}"[:80]
    folder_path = Path(output_dir) / folder_name
    folder_path.mkdir(parents=True, exist_ok=True)

    # Derive candidate name for file
    name = rewritten_resume.get("name", "Resume").replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_filename = f"{name}_{slugify(company)}_{slugify(role)}.docx"
    out_path = folder_path / out_filename

    # Copy original to output (we edit the copy)
    shutil.copy2(original_docx_path, out_path)
    print(f"[Patcher] Copied original DOCX to: {out_path}")

    # Open the copy for editing
    doc = Document(str(out_path))

    # Collect all paragraphs (body + tables)
    all_paras = list(_get_all_paragraphs_from_doc(doc))
    print(f"[Patcher] Found {len(all_paras)} paragraphs in original DOCX")

    # ── Build lookup of old bullet text from original → new bullet text ──────
    # We match by: for each rewritten role, for each new bullet,
    # find the paragraph in the original that is the closest match and replace it.

    rewritten_experience = rewritten_resume.get("experience", [])

    # Collect all (para_text, para_object) pairs that look like bullet lines
    # (non-empty, not headers, length between 20–500 chars)
    bullet_paras = [
        (p, _get_para_full_text(p))
        for p in all_paras
        if 20 <= len(_get_para_full_text(p).strip()) <= 500
    ]

    matched_indices = set()
    replacements_made = 0

    for exp in rewritten_experience:
        new_bullets = exp.get("bullets", [])
        old_title = exp.get("title", "")
        old_company = exp.get("company", "")

        # Find the anchor: the title paragraph for this role
        # This lets us do replacement only within the right section
        anchor_score_threshold = 0.5
        anchor_idx = None

        for i, (para, text) in enumerate(bullet_paras):
            title_score = _fuzzy_match_score(old_title, text)
            company_score = _fuzzy_match_score(old_company, text)
            if title_score >= anchor_score_threshold or company_score >= anchor_score_threshold:
                anchor_idx = i
                break

        # Window: search within reasonable range after the role header
        if anchor_idx is not None:
            search_start = anchor_idx
            search_end = min(anchor_idx + 25, len(bullet_paras))
        else:
            search_start = 0
            search_end = len(bullet_paras)

        # Now match each new bullet to the closest original bullet in range
        for new_bullet in new_bullets:
            best_score = 0.35  # minimum threshold to replace
            best_idx = None

            for i in range(search_start, search_end):
                if i in matched_indices:
                    continue
                _, orig_text = bullet_paras[i]
                score = _fuzzy_match_score(new_bullet, orig_text)
                if score > best_score:
                    best_score = score
                    best_idx = i

            if best_idx is not None:
                para_obj, old_text = bullet_paras[best_idx]
                print(f"[Patcher] Replacing (score={best_score:.2f}):")
                print(f"          OLD: {old_text[:80]}...")
                print(f"          NEW: {new_bullet[:80]}...")
                _set_para_text_preserve_format(para_obj, new_bullet)
                matched_indices.add(best_idx)
                replacements_made += 1

    print(f"[Patcher] Total replacements: {replacements_made}")

    # Also update summary if present
    rewritten_summary = rewritten_resume.get("summary", "")
    if rewritten_summary:
        for para, text in bullet_paras:
            score = _fuzzy_match_score(rewritten_summary[:80], text[:80])
            if score >= 0.45:
                _set_para_text_preserve_format(para, rewritten_summary)
                print(f"[Patcher] Updated summary paragraph")
                break

    # Save patched document
    doc.save(str(out_path))
    print(f"[Patcher] Saved patched DOCX: {out_path}")

    return str(out_path)


def get_original_docx_path() -> str | None:
    """
    Returns the path to the user's original uploaded DOCX, stored as
    'master_resume_original.docx' alongside base_resume.json.
    Returns None if not found.
    """
    base_dir = Path(__file__).parent
    candidates = [
        base_dir / "master_resume_original.docx",
        base_dir / "master_resume_original.pdf",
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return None
